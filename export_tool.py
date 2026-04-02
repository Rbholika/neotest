"""
tools/export_tool.py — PDF, DOCX & Markdown Export
====================================================
Standalone tool — no agent dependency required.

Usage:
    from tools.export_tool import ExportTool
    tool = ExportTool()
    path = tool.to_pdf({"title": "Report", "response": "..."})
    path = tool.to_docx({"title": "Report", "response": "..."})
    path = tool.to_markdown({"title": "Report", "response": "..."})
"""

import os
import json
from datetime import datetime
from typing import Dict, Any, Optional

from config import CONFIG
from core.logger import get_logger

logger = get_logger("tool.export")


class ExportTool:

    def __init__(self, export_dir: str = None):
        self.export_dir = export_dir or CONFIG.EXPORT_DIR
        os.makedirs(self.export_dir, exist_ok=True)

    def _timestamp(self) -> str:
        return datetime.now().strftime("%Y%m%d_%H%M%S")

    def _resolve_path(self, filename: str, ext: str) -> str:
        name = f"{filename}_{self._timestamp()}.{ext}"
        return os.path.join(self.export_dir, name)

    # ── PDF ───────────────────────────────────────────────────────────────────

    def to_pdf(
        self,
        content: Dict[str, Any],
        filename: str = "report",
    ) -> str:
        """
        content keys (all optional):
            title, query, response, sources, sections (list of {heading, body})
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            Table, TableStyle, HRFlowable,
        )

        path = self._resolve_path(filename, "pdf")
        doc  = SimpleDocTemplate(
            path,
            pagesize=A4,
            leftMargin=2*cm, rightMargin=2*cm,
            topMargin=2*cm,  bottomMargin=2*cm,
        )

        styles   = getSampleStyleSheet()
        title_s  = ParagraphStyle("title",  parent=styles["Title"],  fontSize=20, spaceAfter=12)
        h1_s     = ParagraphStyle("h1",     parent=styles["Heading1"], fontSize=14, spaceAfter=6)
        body_s   = ParagraphStyle("body",   parent=styles["Normal"],  fontSize=10, leading=15)
        meta_s   = ParagraphStyle("meta",   parent=styles["Normal"],  fontSize=9,
                                   textColor=colors.grey)

        story = []

        # Title
        story.append(Paragraph(content.get("title", "Report"), title_s))
        story.append(Paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            meta_s,
        ))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.lightgrey))
        story.append(Spacer(1, 0.3*cm))

        # Query (if present)
        if content.get("query"):
            story.append(Paragraph("Query", h1_s))
            story.append(Paragraph(content["query"], body_s))
            story.append(Spacer(1, 0.3*cm))

        # Sources (if present)
        if content.get("sources"):
            story.append(Paragraph("Sources", h1_s))
            for src in content["sources"]:
                story.append(Paragraph(f"• {src}", body_s))
            story.append(Spacer(1, 0.3*cm))

        # Arbitrary sections
        for section in content.get("sections", []):
            story.append(Paragraph(section.get("heading", ""), h1_s))
            story.append(Paragraph(
                str(section.get("body", "")).replace("\n", "<br/>"),
                body_s,
            ))
            story.append(Spacer(1, 0.2*cm))

        # Main response / body
        if content.get("response"):
            story.append(Paragraph("Analysis", h1_s))
            for line in str(content["response"]).split("\n"):
                if line.strip():
                    story.append(Paragraph(line, body_s))
            story.append(Spacer(1, 0.2*cm))

        # Raw data table (if dict/list provided)
        if content.get("table_data") and isinstance(content["table_data"], list):
            story.append(Paragraph("Data", h1_s))
            rows = content["table_data"]
            if rows:
                headers = list(rows[0].keys())
                tbl_data = [headers] + [[str(r.get(h, "")) for h in headers] for r in rows[:50]]
                tbl = Table(tbl_data, repeatRows=1)
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563EB")),
                    ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE",   (0, 0), (-1, -1), 8),
                    ("GRID",       (0, 0), (-1, -1), 0.5, colors.lightgrey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F4F6")]),
                ]))
                story.append(tbl)

        doc.build(story)
        logger.info(f"PDF exported → {path}")
        return path

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def to_docx(
        self,
        content: Dict[str, Any],
        filename: str = "report",
    ) -> str:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        path = self._resolve_path(filename, "docx")
        doc  = Document()

        # Styles
        doc.core_properties.title   = content.get("title", "Report")
        doc.core_properties.author  = "AI Agent Framework"

        # Title
        t = doc.add_heading(content.get("title", "Report"), level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        meta.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        meta.runs[0].font.size = Pt(9)

        doc.add_paragraph()

        # Query
        if content.get("query"):
            doc.add_heading("Query", level=2)
            doc.add_paragraph(content["query"])

        # Sources
        if content.get("sources"):
            doc.add_heading("Sources", level=2)
            for src in content["sources"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(src)

        # Sections
        for section in content.get("sections", []):
            doc.add_heading(section.get("heading", ""), level=2)
            doc.add_paragraph(str(section.get("body", "")))

        # Response
        if content.get("response"):
            doc.add_heading("Analysis", level=2)
            for line in str(content["response"]).split("\n"):
                if line.strip():
                    doc.add_paragraph(line)

        # Table data
        if content.get("table_data") and isinstance(content["table_data"], list):
            rows = content["table_data"][:51]
            if rows:
                doc.add_heading("Data", level=2)
                headers = list(rows[0].keys())
                tbl = doc.add_table(rows=1, cols=len(headers))
                tbl.style = "Light Shading Accent 1"
                hdr = tbl.rows[0].cells
                for i, h in enumerate(headers):
                    hdr[i].text = h
                for row in rows:
                    cells = tbl.add_row().cells
                    for i, h in enumerate(headers):
                        cells[i].text = str(row.get(h, ""))

        doc.save(path)
        logger.info(f"DOCX exported → {path}")
        return path

    # ── Markdown ──────────────────────────────────────────────────────────────

    def to_markdown(
        self,
        content: Dict[str, Any],
        filename: str = "report",
    ) -> str:
        path  = self._resolve_path(filename, "md")
        lines = []

        lines.append(f"# {content.get('title', 'Report')}")
        lines.append(
            f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n"
        )

        if content.get("query"):
            lines.append("## Query")
            lines.append(content["query"] + "\n")

        if content.get("sources"):
            lines.append("## Sources")
            for src in content["sources"]:
                lines.append(f"- {src}")
            lines.append("")

        for section in content.get("sections", []):
            lines.append(f"## {section.get('heading', '')}")
            lines.append(str(section.get("body", "")) + "\n")

        if content.get("response"):
            lines.append("## Analysis")
            lines.append(str(content["response"]) + "\n")

        if content.get("table_data") and isinstance(content["table_data"], list):
            rows = content["table_data"]
            if rows:
                lines.append("## Data")
                headers = list(rows[0].keys())
                lines.append("| " + " | ".join(headers) + " |")
                lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                for row in rows[:50]:
                    lines.append(
                        "| " + " | ".join(str(row.get(h, "")) for h in headers) + " |"
                    )
                lines.append("")

        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

        logger.info(f"Markdown exported → {path}")
        return path
